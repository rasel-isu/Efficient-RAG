import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#000000"},
            bottom={"sz": 12, "val": "single", "color": "#000000"},
            start={"sz": 12, "val": "single", "color": "#000000"},
            end={"sz": 12, "val": "single", "color": "#000000"},
        )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcPr.append(element)


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    """Merge cells horizontally in a row"""
    cell = table.rows[row_idx].cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell.merge(table.rows[row_idx].cells[col])
    return cell


def merge_cells_vertical(table, col_idx, start_row, end_row):
    """Merge cells vertically in a column"""
    cell = table.rows[start_row].cells[col_idx]
    for row in range(start_row + 1, end_row + 1):
        cell.merge(table.rows[row].cells[col_idx])
    return cell


def format_header_cell(cell, text, bg_color, font_size=11, bold=True):
    """Format a header cell with text, background color, and styling"""
    set_cell_background(cell, bg_color)
    set_cell_border(
        cell,
        top={"sz": 4, "val": "single", "color": "#000000"},
        bottom={"sz": 4, "val": "single", "color": "#000000"},
        start={"sz": 4, "val": "single", "color": "#000000"},
        end={"sz": 4, "val": "single", "color": "#000000"},
    )
    
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Arial'


def format_data_cell(cell, text, bg_color="FFFFFF", centered=True):
    """Format a data cell"""
    set_cell_background(cell, bg_color)
    set_cell_border(
        cell,
        top={"sz": 4, "val": "single", "color": "#CCCCCC"},
        bottom={"sz": 4, "val": "single", "color": "#CCCCCC"},
        start={"sz": 4, "val": "single", "color": "#CCCCCC"},
        end={"sz": 4, "val": "single", "color": "#CCCCCC"},
    )
    
    paragraph = cell.paragraphs[0]
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.font.size = Pt(9)
    run.font.name = 'Arial'


def get_value(obj, path, default="N/A"):
    """Safely get nested dictionary value"""
    keys = path.split('.')
    value = obj
    for key in keys:
        if isinstance(value, dict) and key in value:
            value = value[key]
        else:
            return default
        

    return value if value is not None else default


def format_number(value):
    """Format number to 4 decimal places"""
    if isinstance(value, int):
        return f"{value}"
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def load_experiments(file_paths):
    """Load experiment JSON files"""
    experiments = []
    for file_path in file_paths:
        try:
            with open(file_path, 'r') as f:
                data = json.load(f)
                experiments.append(data)
                print(f"✅ Loaded: {file_path}")
        except Exception as e:
            print(f"❌ Error loading {file_path}: {e}")
    return experiments


def create_main_metrics_table(doc, experiments):
    """Create main metrics comparison table"""
    
    # Add section heading
    heading = doc.add_heading('Overall Metrics Comparison', level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Create table: 2 header rows + data rows, 13 columns
    num_rows = 2 + len(experiments)
    num_cols = 13
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Set column widths (approximate, in inches)
    col_widths = [1.0, 0.6, 0.6, 0.6, 0.5, 0.5, 0.5, 0.5, 0.5, 0.5, 0.6, 0.7, 0.7]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # First header row
    row0 = table.rows[0]
    
    # Column 0: Experiment Name (spans 2 rows)
    format_header_cell(row0.cells[0], "Experiment\nName", "4472C4", font_size=10)
    merge_cells_vertical(table, 0, 0, 1)
    
    # Columns 1-3: Simple headers (span 2 rows)
    format_header_cell(row0.cells[1], "Exact\nMatch", "4472C4", font_size=9)
    merge_cells_vertical(table, 1, 0, 1)
    
    format_header_cell(row0.cells[2], "Partial\nMatch", "4472C4", font_size=9)
    merge_cells_vertical(table, 2, 0, 1)
    
    format_header_cell(row0.cells[3], "Mean\nF1", "4472C4", font_size=9)
    merge_cells_vertical(table, 3, 0, 1)
    
    # Columns 4-6: BERT Score (merged header)
    format_header_cell(row0.cells[4], "BERT Score", "4472C4")
    merge_cells_horizontal(table, 0, 4, 6)
    
    # Columns 7-9: ROUGE Scores (merged header)
    format_header_cell(row0.cells[7], "ROUGE Scores", "4472C4")
    merge_cells_horizontal(table, 0, 7, 9)
    
    # Columns 10-12: Simple headers (span 2 rows)
    format_header_cell(row0.cells[10], "Semantic\nSim", "4472C4", font_size=9)
    merge_cells_vertical(table, 10, 0, 1)
    
    format_header_cell(row0.cells[11], "Avg Prompt\nTokens", "4472C4", font_size=9)
    merge_cells_vertical(table, 11, 0, 1)
    
    format_header_cell(row0.cells[12], "Total Prompt\nTokens", "4472C4", font_size=9)
    merge_cells_vertical(table, 12, 0, 1)
    
    # Second header row (sub-columns)
    row1 = table.rows[1]
    
    # BERT sub-columns
    format_header_cell(row1.cells[4], "Precision", "5B9BD5", font_size=9)
    format_header_cell(row1.cells[5], "Recall", "5B9BD5", font_size=9)
    format_header_cell(row1.cells[6], "F1", "5B9BD5", font_size=9)
    
    # ROUGE sub-columns
    format_header_cell(row1.cells[7], "ROUGE-1", "5B9BD5", font_size=9)
    format_header_cell(row1.cells[8], "ROUGE-2", "5B9BD5", font_size=9)
    format_header_cell(row1.cells[9], "ROUGE-L", "5B9BD5", font_size=9)
    
    # Data rows
    for idx, exp in enumerate(experiments):
        row = table.rows[idx + 2]
        bg_color = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
        
        # Column data with paths
        data_columns = [
            (get_value(exp, 'experiment_name', 'Unknown'), False),
            (format_number(get_value(exp, 'basic_metrics.exact_match_accuracy')), True),
            (format_number(get_value(exp, 'basic_metrics.partial_match_accuracy')), True),
            (format_number(get_value(exp, 'basic_metrics.mean_f1_score')), True),
            (format_number(get_value(exp, 'bert_score.precision')), True),
            (format_number(get_value(exp, 'bert_score.recall')), True),
            (format_number(get_value(exp, 'bert_score.f1')), True),
            (format_number(get_value(exp, 'rouge_scores.rouge1.mean')), True),
            (format_number(get_value(exp, 'rouge_scores.rouge2.mean')), True),
            (format_number(get_value(exp, 'rouge_scores.rougeL.mean')), True),
            (format_number(get_value(exp, 'semantic_similarity.mean')), True),
            (format_number(get_value(exp, 'token_statistics.avg_prompt_tokens')), True),
            (format_number(get_value(exp, 'token_statistics.total_prompt_tokens')), True),
        ]
        
        for col_idx, (value, centered) in enumerate(data_columns):
            format_data_cell(row.cells[col_idx], value, bg_color, centered)


def create_question_type_table(doc, experiments):
    """Create by_question_type breakdown table"""
    
    # Add spacing
    doc.add_paragraph()
    
    # Add section heading
    heading = doc.add_heading('Performance by Question Type', level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Question types to include (in order)
    question_types = ['yes_no', 'what', 'how', 'when', 'who', 'where', 'which', 'why', 'other']
    
    # Create table: 2 header rows + (question_types * experiments) data rows, 5 columns
    # Columns: Experiment Name | Question Type | Count | Exact Match Accuracy | Mean F1 Score
    num_data_rows = len(experiments) * len(question_types)
    num_rows = 2 + num_data_rows
    num_cols = 5
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Set column widths
    col_widths = [2.0, 1.2, 0.8, 1.3, 1.3]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
    
    # First header row
    row0 = table.rows[0]
    
    # Experiment Name (spans 2 rows)
    format_header_cell(row0.cells[0], "Experiment Name", "4472C4", font_size=10)
    merge_cells_vertical(table, 0, 0, 1)
    
    # Question Type (spans 2 rows)
    format_header_cell(row0.cells[1], "Question Type", "4472C4", font_size=10)
    merge_cells_vertical(table, 1, 0, 1)
    
    # Count (spans 2 rows)
    format_header_cell(row0.cells[2], "Count", "4472C4", font_size=10)
    merge_cells_vertical(table, 2, 0, 1)
    
    # Performance Metrics (spans 2 columns)
    format_header_cell(row0.cells[3], "Performance Metrics", "4472C4", font_size=10)
    merge_cells_horizontal(table, 0, 3, 4)
    
    # Second header row (sub-columns)
    row1 = table.rows[1]
    
    format_header_cell(row1.cells[3], "Exact Match Accuracy", "5B9BD5", font_size=9)
    format_header_cell(row1.cells[4], "Mean F1 Score", "5B9BD5", font_size=9)
    
    # Data rows
    row_idx = 2
    for exp_idx, exp in enumerate(experiments):
        exp_name = get_value(exp, 'experiment_name', 'Unknown')
        by_question_type = get_value(exp, 'by_question_type', {})
        
        for qt_idx, question_type in enumerate(question_types):
            row = table.rows[row_idx]
            bg_color = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
            
            # Get question type data
            qt_data = by_question_type.get(question_type, {})
            count = get_value(qt_data, 'count', 'N/A')
            exact_match = format_number(get_value(qt_data, 'exact_match_accuracy', 'N/A'))
            mean_f1 = format_number(get_value(qt_data, 'mean_f1_score', 'N/A'))
            
            # Format question type name
            qt_display = question_type.replace('_', ' ').title()
            
            # Column data
            data_columns = [
                (exp_name if qt_idx == 0 else "", False),  # Show experiment name only on first row
                (qt_display, False),
                (str(count), True),
                (exact_match, True),
                (mean_f1, True),
            ]
            
            for col_idx, (value, centered) in enumerate(data_columns):
                format_data_cell(row.cells[col_idx], value, bg_color, centered)
            
            row_idx += 1
        
        # Merge experiment name cells vertically for this experiment
        if len(question_types) > 1:
            start_row = 2 + (exp_idx * len(question_types))
            end_row = start_row + len(question_types) - 1
            merge_cells_vertical(table, 0, start_row, end_row)


def create_experiment_tables(experiments_name, experiments, output_file="experiment_results.docx"):
    """Create Word document with experiment comparison tables"""
    
    # Create document
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.page_height = Inches(8.5)
    section.page_width = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add main title
    title = doc.add_heading(experiments_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add main metrics table
    create_main_metrics_table(doc, experiments)
    
    # Add page break
    doc.add_page_break()
    
    # Add question type breakdown table
    create_question_type_table(doc, experiments)
    
    # Save document
    doc.save(output_file)
    print(f"\n✅ Document created: {output_file}")
    print(f"   Experiments included: {len(experiments)}")
    print(f"   Tables: 2 (Main Metrics + Question Type Breakdown)")


def main():
    # if len(sys.argv) < 2:
    #     print("Usage: python create_experiment_docx.py <experiment1.json> <experiment2.json> ...")
    #     print("\nExample:")
    #     print("  python create_experiment_docx.py experiment1.json experiment2.json")
    #     print("  python create_experiment_docx.py results/*.json")
    #     sys.exit(1)
    
    # Load experiments
    # json_files = sys.argv[1:]

    # experiments_name = 'Llama-3.2-3B-Instruct on rag-mini-wikipedia'
    # base_dir = 'OUTPUT/rag-mini-wikipedia/Llama-3.2-3B-Instruct/'

    experiments_name = 'gpt-3.5-turbo on rag-mini-wikipedia'
    base_dir = 'OUTPUT/rag-mini-wikipedia/gpt-3.5-turbo/'

    json_files = [
        f'{base_dir}baseline_rag_performence.json',
        f'{base_dir}t5_large_sumry_performence.json',
        f'{base_dir}t5_base_sumry_performence.json',
        f'{base_dir}t5_small_sumry_performence.json',
        f'{base_dir}no_keyword_based_filtering_only_sumry_performence.json',
        f'{base_dir}no_sumry_only_keyword_based_filtering_performence.json',

    ]
    experiments = load_experiments(json_files)
    
    if not experiments:
        print("No valid experiments loaded")
        sys.exit(1)
    
    
    output_file = f"{base_dir}experiment_results.docx"
    create_experiment_tables(experiments_name, experiments, output_file)


if __name__ == "__main__":
    main()